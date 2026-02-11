#!/usr/bin/env python3
"""
Generate .docx files for philosophical reflections
Uses python-docx if available, otherwise creates markdown that can be converted
"""

import sys
from pathlib import Path
from datetime import datetime

def create_philosophical_docx(title, author, date, content, output_path):
    """Create a formatted document for philosophical reflections"""
    
    # Try to use python-docx
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title, 0)
        
        # Author and metadata
        doc.add_paragraph(f"von {author}")
        doc.add_paragraph(f"Reflexion vom {date}")
        doc.add_paragraph()
        
        # Content
        doc.add_heading("Zusammenfassung und Reflexion", level=1)
        
        # Split content into paragraphs and add
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("—").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Philosophische Bibliothek 100 | Tägliche Reflexion").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_path)
        print(f"✓ Dokument erstellt: {output_path}")
        return True
        
    except ImportError:
        # Fallback: Create markdown that user can convert
        md_content = f"""# {title}

**von {author}**  
*Reflexion vom {date}*

---

## Zusammenfassung und Reflexion

{content}

---

*Philosophische Bibliothek 100 | Tägliche Reflexion*
"""
        output_md = output_path.replace('.docx', '.md')
        with open(output_md, 'w') as f:
            f.write(md_content)
        print(f"⚠ python-docx nicht verfügbar, Markdown erstellt: {output_md}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 5:
        print("Usage: python create_philo_doc.py 'Title' 'Author' 'Date' 'Content' output.docx")
        sys.exit(1)
    
    title = sys.argv[1]
    author = sys.argv[2]
    date = sys.argv[3]
    content = sys.argv[4]
    output = sys.argv[5]
    
    create_philosophical_docx(title, author, date, content, output)
